#!/usr/bin/env python3
"""Generate a comprehensive Word document wiki for SQLancer test oracles."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global style setup ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helper functions ──

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_cmd(doc, cmd_text):
    p = doc.add_paragraph()
    run = p.add_run(cmd_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_note(doc, text):
    p = doc.add_paragraph()
    run_label = p.add_run('[注意] ')
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run_label.font.size = Pt(10)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_bold_text(doc, label, text):
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(10.5)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════
# SECTION 1: 文档封面与目录引导
# ══════════════════════════════════════════════════════════════════

doc.add_paragraph()
title = doc.add_heading('SQLancer 新增 Test Oracle 全景指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('功能 · 算法 · 论文 · 覆盖 · 场景 · 命令')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('面向：数据库测试骨干 / 集成测试人员 / 开发工程师\n版本：v2.0.61 | 2026-05-28')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# 目录引导
doc.add_heading('目录', level=2)
toc_items = [
    '1  全景概览：Oracle 分类与支持矩阵',
    '2  表达式等价类 Oracle — EET 系列',
    '3  查询计划类 Oracle — DQP / CERT / QPG',
    '4  DML 等价类 Oracle — DQE',
    '5  数据库构造类 Oracle — EDC',
    '6  表达式折叠类 Oracle — CODDTEST',
    '7  优化等价类 Oracle — SONAR',
    '8  事务隔离类 Oracle — WRITE_CHECK / TX_INFER / FUCCI',
    '9  附录：全局参数速查表',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1: 全景概览
# ══════════════════════════════════════════════════════════════════

doc.add_heading('1  全景概览：Oracle 分类与支持矩阵', level=1)

doc.add_paragraph(
    'SQLancer 通过"差异测试"（differential testing）思想检测数据库管理系统（DBMS）中的逻辑错误。'
    '每类 Oracle 的核心思路相同：生成语义等价的两个查询/操作，执行后比对结果；结果不一致则标记为潜在 Bug。'
    '本文档覆盖所有新增 Oracle（超出上游 SQLancer 的扩展），包括其核心算法、参考论文、语法覆盖范围、适用场景和具体命令。'
)

doc.add_heading('1.1  Oracle 分类体系', level=2)

add_table(doc,
    ['分类', 'Oracle', '检测原理简述', '参考论文'],
    [
        ['表达式等价', 'EET / EET_UPDATE / EET_DELETE / EET_INSERT_SELECT',
         '等价表达式变换：变换 WHERE/SELECT 表达式后比对查询结果', 'EET 原生工具 (OSDI 2024)'],
        ['查询计划', 'DQP', '强制不同执行计划，比对同一逻辑查询的结果', 'SIGMOD 2024'],
        ['查询计划', 'CERT', '查询微调后比对估计行数与计划一致性', 'ICSE 2024'],
        ['查询计划', 'QPG', '变异数据库状态以探索未见查询计划', '实验性 (无论文)'],
        ['DML 等价', 'DQE', '同一 WHERE 条件下 SELECT/UPDATE/DELETE 应访问相同行集', 'IEEE 2023'],
        ['数据库构造', 'EDC', '有约束 vs 无约束数据库上同一查询应返回相同结果', 'RADAR (ISSTA 2024)'],
        ['表达式折叠', 'CODDTEST', '常量折叠：表达式替换为预计算常量后结果应一致', 'SIGMOD 2025'],
        ['优化等价', 'SONAR', '优化路径 vs 非优化路径查询结果应一致', 'SemBug (ISSTA 2024)'],
        ['事务隔离', 'WRITE_CHECK', '并发调度下事务结果与等价 Oracle 调度比对', 'WriteCheck 工具 (无正式论文)'],
        ['事务隔离', 'TX_INFER', 'MVCC 版本推理：用辅助版本表推断预期结果', '内部扩展'],
        ['事务隔离', 'FUCCI (DT/MT/CS)', '组合差分测试/蜕变测试/约束求解检测隔离 Bug', 'Fucci (ACM 2024)'],
    ],
    col_widths=[3, 6, 8, 5]
)

doc.add_heading('1.2  支持矩阵', level=2)

add_table(doc,
    ['Oracle', 'MySQL', 'PostgreSQL', 'GaussDB-M', 'GaussDB-A', 'SQLite3', 'CockroachDB', 'TiDB', 'MariaDB'],
    [
        ['EET', 'Y', 'Y', 'Y', 'Y', '-', '-', '-', '-'],
        ['EET_UPDATE', 'Y', 'Y', 'Y', 'Y', '-', '-', '-', '-'],
        ['EET_DELETE', 'Y', 'Y', 'Y', 'Y', '-', '-', '-', '-'],
        ['EET_INSERT_SELECT', 'Y', 'Y', 'Y', 'Y', '-', '-', '-', '-'],
        ['DQP', 'Y', 'Y', 'Y', 'Y', '-', '-', 'Y', 'Y'],
        ['CERT', 'Y', 'Y', 'Y', 'Y', 'Y', 'Y', '-', '-'],
        ['QPG', 'Y', 'Y', 'Y', 'Y', 'Y', 'Y', '-', '-'],
        ['DQE', 'Y', 'Y', 'Y', 'Y', '-', '-', '-', '-'],
        ['EDC', 'Y', 'Y', 'Y', '-', '-', '-', '-', '-'],
        ['CODDTEST', 'Y', 'Y', 'Y', '-', 'Y', '-', '-', '-'],
        ['SONAR', 'Y', 'Y', 'Y', '-', 'Y', '-', 'Y', 'Y'],
        ['WRITE_CHECK', 'Y', 'Y', 'Y', 'Y', '-', 'Y', '-', '-'],
        ['TX_INFER', 'Y', '-', 'Y', '-', '-', '-', '-', '-'],
        ['FUCCI', 'Y', 'Y', 'Y', 'Y', '-', '-', '-', '-'],
    ],
    col_widths=[4, 2, 2, 2.5, 2.5, 2, 2.5, 1.5, 1.5]
)

add_note(doc, 'GaussDB-A 使用 --target-database 参数指定 Oracle 兼容模式连接的目标数据库；GaussDB-PG 的 Oracle 子集与 PostgreSQL 重合。')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2: EET 系列
# ══════════════════════════════════════════════════════════════════

doc.add_heading('2  表达式等价类 Oracle — EET 系列', level=1)

doc.add_heading('2.1  核心原理', level=2)

doc.add_paragraph(
    'EET（Equivalent Expression Transformation）的核心公式：若表达式 E 与 E\' 语义等价（E ≡ E\'），'
    '则在任意查询 Q 中将 E 替换为 E\' 所得 Q\' 应产生相同结果集：'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('E ≡ E\'  ⇒  DB(Q) ≡ DB(Q\')')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    '当 DBMS 返回不同结果时，说明优化器或执行引擎在处理等价表达式时产生了逻辑错误。'
    'EET 使用 SQL 三值逻辑（TRUE / FALSE / NULL）构造变换，确保变换在所有 NULL 边界条件下仍然语义等价。'
)

add_bold_text(doc, '参考论文：', 'EET 原生工具，源自 OSDI 2024 论文 "Equivalent Expression Transformation for Detecting Logic Bugs in DBMS"。SQLancer 版本在原生工具基础上重构为统一框架，支持多方言 AST 适配。')

doc.add_heading('2.2  变换规则体系', level=2)

doc.add_paragraph('EET 包含两类变换：Wrapping 规则（包裹式）和语义重写规则（结构替换式）。')

doc.add_heading('2.2.1  Wrapping 规则 (Rules 1–6 + ConstBool)', level=3)

add_table(doc,
    ['规则', '变换模板', '说明'],
    [
        ['Rule 1', 'bool_expr → (p OR NOT p OR p IS NULL) AND bool_expr',
         '永真式 AND 包裹：true_expr(p) = p OR NOT p OR p IS NULL'],
        ['Rule 2', 'bool_expr → (p AND NOT p AND p IS NOT NULL) OR bool_expr',
         '永假式 OR 包裹：false_expr(p) = p AND NOT p AND p IS NOT NULL'],
        ['Rule 3', 'expr → CASE WHEN false_expr THEN rand_expr ELSE expr END',
         '冗余分支：条件恒为 FALSE，永远走 ELSE'],
        ['Rule 4', 'expr → CASE WHEN true_expr THEN expr ELSE rand_expr END',
         '冗余分支：条件恒为 TRUE，永远走 THEN'],
        ['Rule 5', 'expr → CASE WHEN rand_bool THEN copy(expr) ELSE expr END',
         '冗余分支：随机布尔条件，THEN/ELSE 均为 expr'],
        ['Rule 6', 'expr → CASE WHEN rand_bool THEN expr ELSE copy(expr) END',
         '冗余分支：与 Rule 5 对称'],
        ['ConstBool', '1 → 1 OR random_bool, 0 → 0 AND random_bool',
         '常量布尔值扩展：已知 1/0 的布尔语义不变'],
    ],
    col_widths=[2.5, 8, 7]
)

doc.add_heading('2.2.2  语义重写规则', level=3)

add_table(doc,
    ['规则名', '变换', '适用方言'],
    [
        ['DeMorgan', '(A AND B) → NOT(NOT(A) OR NOT(B))', '所有方言'],
        ['BetweenToComparison', 'x BETWEEN a AND b → (x>=a AND x<=b)', '所有方言'],
        ['ExistsToIn', 'EXISTS(Q WHERE p) → TRUE IN (SELECT CASE...WHERE TRUE)', 'MySQL/PG/GaussDB-M/GaussDB-A'],
        ['InToExists', 'lhs IN Q → CASE WHEN IS NOT NULL THEN EXISTS(...) ELSE NULL END', 'MySQL/PG/GaussDB-M/GaussDB-A'],
        ['IntersectToExists', 'Q1 INTERSECT Q2 → Q1 WHERE EXISTS(Q2 WHERE nullSafeEq)', 'PostgreSQL'],
        ['ExceptToNotExists', 'Q1 EXCEPT Q2 → Q1 WHERE NOT EXISTS(Q2 WHERE nullSafeEq)', 'PostgreSQL / GaussDB-A (MINUS)'],
        ['CoalesceToCase', 'COALESCE(a,b) → CASE WHEN a IS NOT NULL THEN a ELSE b END', 'MySQL / GaussDB-M'],
        ['NullifToCase', 'NULLIF(a,b) → CASE WHEN a=b THEN NULL ELSE a END', 'MySQL / GaussDB-M'],
    ],
    col_widths=[3.5, 7.5, 5]
)

add_note(doc, 'nullSafeEq = (Q1.col = Q2.col) OR (Q1.col IS NULL AND Q2.col IS NULL)，逐列 AND 连接。'
         'IN→EXISTS 对 UNION 子查询有专门处理：对每个 UNION 分支注入 lhs=fetch[0] 等值条件。')

doc.add_heading('2.2.3  组件缩减（Component Reduction）', level=3)

doc.add_paragraph(
    '当 EET 变换产生语法错误或运行时异常时，启动组件缩减流程以获得最小可复现查询：'
)

add_table(doc,
    ['阶段', '操作', '目标'],
    [
        ['Phase 1', 'DFS 逆序逐节点替换为 NULL，贪心删除不影响 Bug 触发的节点', '最小化表达式树'],
        ['Phase 2', '逐个剥离 CASE WHEN，替换为 THEN/ELSE 分支', '消除冗余包裹层'],
    ],
    col_widths=[2, 7, 5]
)

doc.add_heading('2.3  语法覆盖范围', level=2)

add_table(doc,
    ['覆盖类别', '具体语法/算子', '方言'],
    [
        ['布尔运算', 'AND, OR, NOT, IS NULL, IS NOT NULL', '全部'],
        ['比较运算', '=, !=, <, <=, >, >=, BETWEEN, IN, EXISTS', '全部'],
        ['算术运算', '+, -, *, /, %', '全部'],
        ['类型转换', 'CAST(expr AS type)', '全部'],
        ['字符串匹配', 'LIKE, SIMILAR TO, POSIX 正则 (~, ~*, !~, !~*)', 'PG'],
        ['集合操作', 'UNION, UNION ALL, INTERSECT, EXCEPT/MINUS', 'PG/GaussDB-A'],
        ['CTE', 'WITH ... AS (...) SELECT', '全部'],
        ['聚合函数', 'COUNT, SUM, AVG, MAX, MIN 等', '全部'],
        ['窗口函数', 'OVER (PARTITION BY ... ORDER BY ...)', 'PG/MySQL'],
        ['子查询', 'EXISTS, IN, ANY/ALL, Scalar Subquery, Lateral', '全部'],
        ['CASE WHEN', 'CASE WHEN ... THEN ... ELSE ... END', '全部'],
        ['DML 语句', 'UPDATE SET, DELETE FROM, INSERT INTO ... SELECT', 'EET_UPDATE/DELETE/INSERT_SELECT'],
    ],
    col_widths=[3, 8, 4]
)

doc.add_heading('2.4  适用场景', level=2)

add_table(doc,
    ['Oracle', '检测目标', '典型 Bug 类型'],
    [
        ['EET', 'SELECT 查询中表达式等价性', '优化器错误变换、NULL 处理错误、常量折叠 Bug'],
        ['EET_UPDATE', 'UPDATE WHERE 条件等价性', 'UPDATE 路径优化 Bug、行选择不一致'],
        ['EET_DELETE', 'DELETE WHERE 条件等价性', 'DELETE 路径优化 Bug、索引选择不一致'],
        ['EET_INSERT_SELECT', 'INSERT...SELECT WHERE 等价性', 'INSERT 子查询优化 Bug'],
    ],
    col_widths=[4, 5, 8]
)

doc.add_heading('2.5  使用命令', level=2)

doc.add_paragraph('SELECT 查询 EET：')
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle EET --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle EET --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle EET --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-a --oracle EET --target-database my_db --num-queries 100')

doc.add_paragraph('DML 操作 EET：')
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle EET_UPDATE --num-queries 50')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle EET_DELETE --num-queries 50')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle EET_INSERT_SELECT --num-queries 50')

add_note(doc, '--target-database 是 GaussDB-A 专属参数（必须指定），用于 Oracle 兼容模式下连接目标数据库。')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3: 查询计划类 — DQP / CERT / QPG
# ══════════════════════════════════════════════════════════════════

doc.add_heading('3  查询计划类 Oracle — DQP / CERT / QPG', level=1)

# ── 3.1 DQP ──

doc.add_heading('3.1  DQP — 确定性查询分区', level=2)

add_bold_text(doc, '参考论文：', '"Keep It Simple: Testing Databases via Differential Query Plans"，SIGMOD 2024。')

doc.add_heading('3.1.1  算法', level=3)

doc.add_paragraph(
    'DQP 利用同一逻辑查询在不同执行计划下应返回相同结果的原理。'
    '通过强制查询优化器选择不同执行路径（optimizer hints、SET 变量），'
    '比对同一 SQL 在不同计划下的结果集。如果结果不一致，则说明优化器在某条路径上产生了逻辑错误。'
)

add_table(doc,
    ['步骤', '操作'],
    [
        ['1', '生成随机 SELECT 查询（含 JOIN、WHERE、GROUP BY、HAVING）'],
        ['2', '执行原始查询，捕获结果集 R0'],
        ['3', '通过 Hint 或 SET 变量强制不同执行计划（如 Index Scan vs Seq Scan）'],
        ['4', '在每种执行计划下执行查询，捕获结果集 R1, R2, ...'],
        ['5', '比对所有结果集：R0 ≡ R1 ≡ R2 ...；不一致则标记 Bug'],
    ],
    col_widths=[1.5, 14]
)

doc.add_heading('3.1.2  语法覆盖', level=3)
doc.add_paragraph('SELECT（含 JOIN、WHERE、GROUP BY、HAVING、DISTINCT、LIMIT）、Optimizer Hints、SET 变量切换执行计划。')

doc.add_heading('3.1.3  适用场景', level=3)
doc.add_paragraph('检测优化器因执行路径选择差异导致的逻辑错误：索引优化 Bug、Hash Join vs Nested Loop 不一致、子查询优化错误。')

doc.add_heading('3.1.4  命令', level=3)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle DQP --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle DQP --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle DQP --num-queries 100')

# ── 3.2 CERT ──

doc.add_heading('3.2  CERT — 基数估计回归测试', level=2)

add_bold_text(doc, '参考论文：', '"CERT: Finding Performance Issues in Database Systems Through the Lens of Cardinality Estimation"，ICSE 2024。')

doc.add_heading('3.2.1  算法', level=3)

doc.add_paragraph(
    'CERT 通过检查查询微调后估计行数（cardinality）的合理性来检测基数估计错误。'
    '核心前提：当两个查询的计划相似（edit distance ≤ 1）时，估计行数的变化方向应与微调方向一致。'
)

add_table(doc,
    ['步骤', '操作'],
    [
        ['1', '生成随机 SELECT 查询 Q0，EXPLAIN 获取估计行数 C0 和计划 P0'],
        ['2', '对 Q0 应用一种微调（加 DISTINCT / 加 WHERE / 加 LIMIT 等），得到 Q1'],
        ['3', 'EXPLAIN Q1 获取 C1 和 P1'],
        ['4', '检查 edit_distance(P0, P1) ≤ 1；否则跳过'],
        ['5', '验证 C0 与 C1 的变化方向与微调预期一致；不一致则标记 Bug'],
    ],
    col_widths=[1.5, 14]
)

doc.add_paragraph('微调类型（8 种）：JOIN 修改、DISTINCT 增删、WHERE 条件增删、GROUP BY 修改、HAVING 修改、AND/OR 替换、LIMIT 修改。')

doc.add_heading('3.2.2  适用场景', level=3)
doc.add_paragraph('基数估计错误、统计信息不准确、查询代价估算 Bug、性能回归问题。注意：CERT 检测的是性能/估计错误而非结果正确性错误。')

doc.add_heading('3.2.3  命令', level=3)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle CERT --num-queries 200')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle CERT --num-queries 200')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle CERT --num-queries 200')

# ── 3.3 QPG ──

doc.add_heading('3.3  QPG — 查询计划引导（实验性）', level=2)

add_bold_text(doc, '参考论文：', '实验性功能，无正式论文。基于强化学习奖励机制探索未见查询计划。')

doc.add_heading('3.3.1  算法', level=3)

doc.add_paragraph(
    'QPG 通过变异数据库状态（增删索引、修改数据分布、ANALYZE 等）来触发优化器选择新的执行计划。'
    '使用加权平均奖励跟踪计划多样性：出现新计划时奖励高，无新计划时触发变异。'
)

add_table(doc,
    ['变异动作', '示例', '触发新计划的原理'],
    [
        ['INSERT', 'INSERT INTO t0(c0,c1) VALUES(\'a\', NULL)', '改变数据分布/空值比例'],
        ['UPDATE', 'UPDATE t0 SET c2=-944 WHERE c1', '改变列值模式'],
        ['DELETE', 'DELETE FROM t0 WHERE c0>3', '减少行数可能改变 JOIN 策略'],
        ['CREATE_INDEX', 'CREATE INDEX i0 ON t0(c0)', '启用 Index Scan'],
        ['DROP_INDEX', 'DROP INDEX i0', '强制 Seq Scan'],
        ['ALTER_TABLE', 'ALTER TABLE t0 ADD COLUMN c39 REAL', 'Schema 变化'],
        ['ANALYZE_TABLE', 'ANALYZE t0', '更新统计信息'],
        ['TRUNCATE', 'TRUNCATE t0', '清空数据重置计划'],
    ],
    col_widths=[3, 6, 7]
)

doc.add_heading('3.3.2  QPG 专属参数', level=3)

add_table(doc,
    ['参数', '默认值', '说明'],
    [
        ['--qpg-enable', 'false', '启用 QPG 功能（必须开启才生效）'],
        ['--qpg-log-query-plan', 'false', '记录每次查询的计划（需先启用 --qpg-enable）'],
        ['--qpg-max-interval', '1000', '无新计划时最多迭代多少次后触发变异'],
        ['--qpg-reward-weight', '0.25', '奖励更新权重 (0-1)，越低对新奖励越敏感'],
        ['--qpg-selection-probability', '0.7', '随机变异概率 (>0.5 偏向探索，<0.5 偏向利用)'],
    ],
    col_widths=[5, 2.5, 8]
)

add_note(doc, '以上 5 个参数为 QPG 专属，仅在 --qpg-enable 开启时有效。')

doc.add_heading('3.3.3  命令', level=3)
add_cmd(doc, 'java -jar sqlancer.jar mysql --qpg-enable --qpg-log-query-plan --num-queries 1000')
add_cmd(doc, 'java -jar sqlancer.jar postgres --qpg-enable --qpg-selection-probability 0.8 --num-queries 500')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-a --qpg-enable --target-database my_db --num-queries 1000')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4: DQE
# ══════════════════════════════════════════════════════════════════

doc.add_heading('4  DML 等价类 Oracle — DQE', level=1)

add_bold_text(doc, '参考论文：', '"Differential Query Execution for Detecting Logic Bugs in Database Systems"，IEEE 2023。')

doc.add_heading('4.1  算法', level=2)

doc.add_paragraph(
    'DQE 的核心前提：同一个 WHERE 条件 φ 在 SELECT、UPDATE、DELETE 三种语句中应访问完全相同的行集。'
    '如果 SELECT 返回的 rowid 集合与 UPDATE 修改的行集或 DELETE 删除的行集不一致，则说明 DBMS 在不同语句路径上对同一谓词的评估不一致。'
)

add_table(doc,
    ['步骤', '操作'],
    [
        ['1', '向目标表添加辅助列：rowid (唯一标识) 和 updated (是否被修改, INT DEFAULT 0)'],
        ['2', '生成随机 WHERE 条件 φ'],
        ['3', '执行 SELECT rowid FROM t WHERE φ → 得到行集 S'],
        ['4', '在 BEGIN/ROLLBACK 中执行 UPDATE t SET updated=1 WHERE φ → 得到行集 U'],
        ['5', '在 BEGIN/ROLLBACK 中执行 DELETE FROM t WHERE φ → 得到行集 D'],
        ['6', '比对 S ≡ U ≡ D；不一致则标记 Bug'],
    ],
    col_widths=[1.5, 14]
)

doc.add_heading('4.2  语法覆盖', level=2)
doc.add_paragraph('WHERE 条件（所有布尔运算、比较运算、IN、EXISTS、LIKE 等）、UPDATE SET、DELETE FROM、事务 BEGIN/ROLLBACK。')

doc.add_heading('4.3  适用场景', level=2)
doc.add_paragraph('SELECT vs UPDATE 行选择不一致、SELECT vs DELETE 行选择不一致、不同语句路径的谓词评估差异、索引在 UPDATE/DELETE 中被错误跳过。')

doc.add_heading('4.4  命令', level=2)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle DQE --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle DQE --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle DQE --num-queries 100')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5: EDC
# ══════════════════════════════════════════════════════════════════

doc.add_heading('5  数据库构造类 Oracle — EDC', level=1)

add_bold_text(doc, '参考论文：', 'RADAR: A Raw Database Construction Framework for DBMS Testing，ISSTA 2024。'
    '作者：Yuyang Rong, Zhiyong Wu, Chengyu Zhang, Manfei Wu, Jiyuan Zhang, Xiangke Liao, Yinfang Chen（国防科技大学）。'
    '工具地址：https://github.com/tcse-iscas/radar')

doc.add_heading('5.1  算法', level=2)

doc.add_paragraph(
    'EDC 构造一个"原始数据库"（Raw DB）：复制所有表结构和数据，但移除所有约束（NOT NULL、UNIQUE、'
    'PRIMARY KEY、FOREIGN KEY、CHECK、GENERATED 列转为普通列）。在有约束的原始 DB 和无约束的 Raw DB '
    '上执行同一查询，如果约束不影响查询语义（仅影响优化路径），结果应相同。'
)

add_table(doc,
    ['步骤', '操作'],
    [
        ['1', '解析 SHOW CREATE TABLE 提取约束信息'],
        ['2', '构造 Raw DB：同表同数据，移除所有约束'],
        ['3', '生成随机 SELECT 查询'],
        ['4', '在原始 DB 和 Raw DB 上分别执行'],
        ['5', '比对结果集；不一致则说明约束优化产生了逻辑错误'],
    ],
    col_widths=[1.5, 14]
)

doc.add_paragraph('多样性控制：使用 hash-based 结构指纹避免重复测试相同约束配置。')

doc.add_heading('5.2  适用场景', level=2)
doc.add_paragraph('NOT NULL 约束优化错误、UNIQUE 约束影响查询结果、FOREIGN KEY 处理 Bug、GENERATED 列计算错误、CHECK 约束被优化器绕过。')

doc.add_heading('5.3  命令', level=2)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle EDC --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle EDC --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle EDC --num-queries 100')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 6: CODDTEST
# ══════════════════════════════════════════════════════════════════

doc.add_heading('6  表达式折叠类 Oracle — CODDTEST', level=1)

add_bold_text(doc, '参考论文：', '"Constant Optimization Driven Database System Testing"，SIGMOD 2025。')

doc.add_heading('6.1  算法', level=2)

doc.add_paragraph(
    'CODDTEST 利用常量折叠（constant folding）原理：如果表达式 E 可以预计算为常量值 C，'
    '则将 E 替换为 C 后查询结果应不变。DBMS 的常量折叠优化如果实现错误，会导致结果偏差。'
)

add_table(doc,
    ['模式 (--coddtest-model)', '测试方式'],
    [
        ['EXPRESSION', '简单 SELECT：WHERE 条件中表达式替换为预计算常量'],
        ['SUBQUERY', '子查询模式：关联子查询中的表达式常量折叠'],
        ['RANDOM', '随机选择 EXPRESSION 或 SUBQUERY 模式（默认）'],
    ],
    col_widths=[5, 11]
)

add_note(doc, '--coddtest-model 是 CODDTEST 专属参数，仅在选择 --oracle CODDTEST 时生效。')

doc.add_heading('6.2  适用场景', level=2)
doc.add_paragraph('常量折叠优化 Bug、表达式预计算错误、关联子查询中常量处理错误、优化器过早折叠导致语义变化。')

doc.add_heading('6.3  命令', level=2)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle CODDTEST --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle CODDTEST --coddtest-model EXPRESSION --num-queries 50')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle CODDTEST --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle CODDTEST --num-queries 100')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 7: SONAR
# ══════════════════════════════════════════════════════════════════

doc.add_heading('7  优化等价类 Oracle — SONAR', level=1)

add_bold_text(doc, '参考论文：', 'SemBug: Detecting Logic Bugs in DBMS via Semantic-Aware Non-Optimizing Query，ISSTA 2024。'
    '作者：Suyang Ju 等。工具地址：https://github.com/Syang111/SemBug。'
    'SONAR 在 SQLancer 中继承了 SemBug 的核心思路：将 WHERE 条件提取为标记列（flag），'
    '通过比对优化路径（直接 WHERE 过滤）与非优化路径（flag 列外层过滤）的结果差异检测优化器错误。'
    '与 NoREC 的区别在于 SONAR 扩展了表达式覆盖范围（窗口函数、聚合函数、60+内置函数等）。')

doc.add_heading('7.1  算法', level=2)

doc.add_paragraph(
    'SONAR 比对"优化路径"和"非优化路径"下同一查询的结果。'
    '优化路径使用 WHERE 条件直接过滤；非优化路径将条件提取为标记列（flag），'
    '在外层查询中通过 WHERE flag=1 过滤。如果两种路径结果不一致，说明优化器变换改变了语义。'
)

add_table(doc,
    ['路径', 'SQL 模式'],
    [
        ['优化', 'SELECT ... FROM t WHERE condition'],
        ['非优化', 'SELECT f1 FROM (SELECT ..., condition IS TRUE AS flag FROM t) sub WHERE flag=1'],
    ],
    col_widths=[2.5, 13.5]
)

doc.add_heading('7.2  适用场景', level=2)
doc.add_paragraph('优化器 WHERE 条件变换 Bug、索引使用导致语义变化、谓词评估顺序错误。')

doc.add_heading('7.3  命令', level=2)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle SONAR --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle SONAR --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle SONAR --num-queries 100')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 8: 事务隔离类 — WRITE_CHECK / TX_INFER / FUCCI
# ══════════════════════════════════════════════════════════════════

doc.add_heading('8  事务隔离类 Oracle', level=1)

doc.add_paragraph(
    '事务隔离类 Oracle 检测并发场景下的隔离级别 Bug，包括脏读、不可重复读、幻读、'
    '写偏序（write skew）等异常。需要多个并发事务和交错调度（schedule）。'
)

doc.add_paragraph(
    '重要说明：WRITE_CHECK 和 FUCCI 是两个独立工具的集成，算法根本不同。'
    'WRITE_CHECK 基于 WriteCheck 工具（调度重排序比对），侧重写可串行化验证；'
    'FUCCI 基于 Fucci 论文（MVCC 模拟 + 约束求解），侧重隔离级别异常检测。'
    '两者互补而非替代。'
)

add_bold_text(doc, '全局事务参数（适用于所有事务隔离类 Oracle）：', '')

add_table(doc,
    ['参数', '默认值', '说明'],
    [
        ['--use-fixed-num-transaction', 'false', '使用固定事务数量（否则随机生成）'],
        ['--num-transaction', '2', '每轮测试的事务数量'],
        ['--num-schedule', '10', '每轮测试生成的调度数量'],
        ['--set-case', 'false', '使用指定测试案例文件（用于 Bug 复现）'],
        ['--case-file', '-', '测试案例文件路径（配合 --set-case）'],
    ],
    col_widths=[5, 2, 8]
)

add_note(doc, '以上 5 个参数为事务类 Oracle 通用参数，不是某个 Oracle 专属。')

# ── 8.1 WRITE_CHECK ──

doc.add_heading('8.1  WRITE_CHECK — 调度比对', level=2)

add_bold_text(doc, '来源工具：', 'WriteCheck — 自动检测写操作可串行化违规的工具。'
    '无正式论文发表。源码位于 D:\\Jack.Xiao\\dbtools\\WriteCheck-main。'
    '已发现 13 个 Bug（1 MySQL, 3 MariaDB, 9 TiDB）。')

doc.add_heading('8.1.1  算法', level=3)

doc.add_paragraph(
    'WRITE_CHECK 生成并发事务和交错调度，通过构造"Oracle 调度"（等价但更可预测的执行序列）'
    '来比对结果。核心检查：事务在遇到死锁时回滚后的行为应与显式 ROLLBACK 一致。'
)

add_table(doc,
    ['步骤', '操作'],
    [
        ['1', '生成多个并发事务（含 BEGIN/SELECT/INSERT/UPDATE/DELETE/COMMIT/ROLLBACK）'],
        ['2', '生成交错调度 S（随机排列各事务的语句顺序）'],
        ['3', '执行 S，捕获每个事务的结果和最终 DB 状态'],
        ['4', '构造 Oracle 调度 S\'：等价执行（死锁点回滚 / 成功事务完整执行）'],
        ['5', '执行 S\'，捕获结果'],
        ['6', '比对 S 和 S\' 的结果；不一致则标记 Bug'],
    ],
    col_widths=[1.5, 14]
)

doc.add_heading('8.1.2  适用场景', level=3)
doc.add_paragraph('死锁处理不一致、隔离级别违反、事务回滚不完全、并发写入丢失。')

doc.add_heading('8.1.3  命令', level=3)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle WRITE_CHECK --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle WRITE_CHECK --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle WRITE_CHECK --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-a --oracle WRITE_CHECK --target-database my_db --num-queries 100')

doc.add_paragraph('Bug 复现：')
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle WRITE_CHECK_REPRODUCE --set-case --case-file bug_case.json')

# ── 8.2 TX_INFER ──

doc.add_heading('8.2  TX_INFER — MVCC 版本推理', level=2)

doc.add_heading('8.2.1  算法', level=3)

doc.add_paragraph(
    'TX_INFER 在数据库中创建辅助版本表（_infer_<table>_vt），追踪每行的版本链（rid, vid, deleted, txid）。'
    '根据隔离级别模拟 MVCC 可见性规则，推断每个事务的预期查询结果，与实际执行结果比对。'
)

add_table(doc,
    ['隔离级别', '可见性规则'],
    [
        ['READ UNCOMMITTED', '可见所有版本（包括未提交）'],
        ['READ COMMITTED', '仅可见已提交版本 + 当前事务版本'],
        ['REPEATABLE READ', '仅可见 BEGIN 时快照中的版本'],
        ['SERIALIZABLE', '全串行化：等价于某种串行执行'],
    ],
    col_widths=[5, 11]
)

doc.add_heading('8.2.2  适用场景', level=3)
doc.add_paragraph('MVCC 实现错误、快照可见性违反、隔离级别语义不符合声明、并发读异常。')

add_note(doc, 'TX_INFER 仅支持 MySQL 和 GaussDB-M，因为需要兼容的 MVCC 实现和辅助表语法。')

doc.add_heading('8.2.3  命令', level=3)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle TX_INFER --num-queries 100')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle TX_INFER --num-queries 100')

# ── 8.3 FUCCI ──

doc.add_heading('8.3  FUCCI — 组合式隔离测试', level=2)

add_bold_text(doc, '参考论文：', 'Fucci: Fuzzing Database Transactions with Random Conflict Construction '
    'and Multilevel Constraint Solving，ACM 2024。DOI: 10.1145/3664102。'
    '作者：Xiang Gao 等。工具源码：D:\\Jack.Xiao\\dbtools\\Fucci-main。')

doc.add_heading('8.3.1  三种子 Oracle', level=3)

add_table(doc,
    ['子 Oracle', '原理', '检测目标'],
    [
        ['DT (差分测试)', '同一调度在不同隔离级别/不同 DBMS 执行，比对结果', '隔离级别间行为不一致'],
        ['MT (蜕变测试)', '等价调度重排列后执行，比对结果', '调度顺序依赖异常'],
        ['CS (约束求解)', '解析 WHERE  predicate 为约束，验证 SELECT 结果满足约束', '谓词评估错误'],
    ],
    col_widths=[3.5, 7, 5.5]
)

doc.add_heading('8.3.2  FUCCI 专属参数', level=3)

add_table(doc,
    ['参数', '默认值', '可选值', '说明'],
    [
        ['--fucci-oracle-type', 'ALL', 'DT/MT/CS/ALL', '选择子 Oracle 类型'],
        ['--fucci-isolation-level', 'RANDOM', 'READ_COMMITTED/REPEATABLE_READ/SERIALIZABLE/RANDOM', '指定隔离级别'],
        ['--fucci-schedule-count', '10', '任意正整数', '每轮测试的调度数量'],
    ],
    col_widths=[4.5, 2.5, 5, 5]
)

add_note(doc, '以上 3 个参数为 FUCCI 专属，仅在 --oracle FUCCI 时生效。')

doc.add_heading('8.3.3  4 层 Bug 缩减', level=3)

add_table(doc,
    ['层级', '缩减策略'],
    [
        ['Layer 1', '语句删除：逐个删除非必要语句，保留触发 Bug 的最小语句集'],
        ['Layer 2', '语句简化：删除列、WHERE 条件等冗余成分'],
        ['Layer 3', '表达式简化：简化 WHERE/SELECT 中的表达式'],
        ['Layer 4', '常量简化：替换复杂常量为简单值'],
    ],
    col_widths=[2.5, 13.5]
)

doc.add_heading('8.3.4  命令', level=3)
add_cmd(doc, 'java -jar sqlancer.jar mysql --oracle FUCCI --fucci-oracle-type DT --fucci-isolation-level SERIALIZABLE')
add_cmd(doc, 'java -jar sqlancer.jar postgres --oracle FUCCI --fucci-oracle-type ALL --fucci-schedule-count 20')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-m --oracle FUCCI --fucci-oracle-type MT')
add_cmd(doc, 'java -jar sqlancer.jar gaussdb-a --oracle FUCCI --fucci-isolation-level REPEATABLE_READ --target-database my_db')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 9: 附录
# ══════════════════════════════════════════════════════════════════

doc.add_heading('9  附录：全局参数速查表', level=1)

doc.add_heading('9.1  通用参数', level=2)

add_table(doc,
    ['参数', '默认值', '适用范围', '说明'],
    [
        ['--num-queries', '1000', '所有 Oracle', '每轮生成的查询数量'],
        ['--timeout', '-', '所有 Oracle', '单条查询超时时间（秒）'],
        ['--username', 'root', '所有 DBMS', '数据库用户名'],
        ['--password', '-', '所有 DBMS', '数据库密码'],
        ['--host', 'localhost', '所有 DBMS', '数据库主机'],
        ['--port', 'DBMS 默认', '所有 DBMS', '数据库端口'],
        ['--database', 'sqlancer', '所有 DBMS', '测试数据库名'],
    ],
    col_widths=[4, 2.5, 3, 7]
)

doc.add_heading('9.2  事务类通用参数', level=2)

add_table(doc,
    ['参数', '默认值', '适用 Oracle', '说明'],
    [
        ['--use-fixed-num-transaction', 'false', 'WRITE_CHECK/TX_INFER/FUCCI', '固定事务数量'],
        ['--num-transaction', '2', 'WRITE_CHECK/TX_INFER/FUCCI', '事务数量'],
        ['--num-schedule', '10', 'WRITE_CHECK/TX_INFER/FUCCI', '调度数量'],
        ['--set-case', 'false', 'WRITE_CHECK_REPRODUCE', '指定测试案例'],
        ['--case-file', '-', 'WRITE_CHECK_REPRODUCE', '案例文件路径'],
    ],
    col_widths=[5, 2, 4, 6]
)

doc.add_heading('9.3  Oracle 专属参数', level=2)

add_table(doc,
    ['参数', '默认值', '所属 Oracle', '说明'],
    [
        ['--qpg-enable', 'false', 'QPG', '启用查询计划引导'],
        ['--qpg-log-query-plan', 'false', 'QPG', '记录查询计划'],
        ['--qpg-max-interval', '1000', 'QPG', '变异触发最大间隔'],
        ['--qpg-reward-weight', '0.25', 'QPG', '奖励更新权重'],
        ['--qpg-selection-probability', '0.7', 'QPG', '随机变异概率'],
        ['--fucci-oracle-type', 'ALL', 'FUCCI', '子 Oracle 类型 (DT/MT/CS/ALL)'],
        ['--fucci-isolation-level', 'RANDOM', 'FUCCI', '隔离级别'],
        ['--fucci-schedule-count', '10', 'FUCCI', '调度数量'],
        ['--target-database', '-', 'GaussDB-A', 'Oracle 兼容模式目标数据库 (必须)'],
        ['--enable-clob-blob', 'false', 'GaussDB-A', '启用 CLOB/BLOB 类型'],
        ['--coddtest-model', 'RANDOM', 'CODDTEST', '测试模式 (EXPRESSION/SUBQUERY/RANDOM)'],
    ],
    col_widths=[5, 2.5, 3, 6]
)

doc.add_heading('9.4  论文引用汇总', level=2)

add_table(doc,
    ['Oracle', '论文', '会议/期刊', '年份'],
    [
        ['EET', 'Equivalent Expression Transformation for Detecting Logic Bugs in DBMS', 'OSDI', '2024'],
        ['DQP', 'Keep It Simple: Testing Databases via Differential Query Plans', 'SIGMOD', '2024'],
        ['CERT', 'CERT: Finding Performance Issues Through Cardinality Estimation', 'ICSE', '2024'],
        ['DQE', 'Differential Query Execution for Detecting Logic Bugs', 'IEEE', '2023'],
        ['CODDTEST', 'Constant Optimization Driven Database System Testing', 'SIGMOD', '2025'],
        ['QPG', '实验性功能', '-', '-'],
        ['EDC', 'RADAR: A Raw Database Construction Framework for DBMS Testing', 'ISSTA', '2024'],
        ['SONAR', 'SemBug: Detecting Logic Bugs via Semantic-Aware Non-Optimizing Query', 'ISSTA', '2024'],
        ['WRITE_CHECK', 'WriteCheck 工具（无正式论文）', '-', '-'],
        ['TX_INFER', '内部扩展', '-', '-'],
        ['FUCCI', 'Fucci: Fuzzing Database Transactions with Random Conflict Construction and Multilevel Constraint Solving', 'ACM', '2024'],
    ],
    col_widths=[3.5, 7, 3, 2]
)

# ── Save ──

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'SQLancer_Test_Oracle_Wiki.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')